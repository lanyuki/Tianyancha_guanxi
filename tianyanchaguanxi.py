from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Inches
import tkinter as tk
import webbrowser
from tkinter import filedialog
import time
import pandas as pd
import random
import os
import sys
import traceback

# 全局变量
excel_file = 'D:\\Tianyancha_guanxi\\11.xlsx'
downloads_folder = 'D:\\Tianyancha_guanxi\\data'

def choose_excel_file():
    global excel_file
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        excel_file = file_path
        excel_label.config(text=f"已选择的 Excel 文件：{excel_file}")

def choose_downloads_folder():
    global downloads_folder
    folder_path = filedialog.askdirectory()
    if folder_path:
        downloads_folder = folder_path
        folder_label.config(text=f"已选择的下载文件夹：{downloads_folder}")

def show_instructions():
    instructions = """本程序通过selenium模拟搜索并批量下载天眼查公司关系图，效果如下：
            1. 需要用户手动登陆，预留时间为30秒。
            2. 用户需要选择其需要打开的 Excel 文件和默认的下载路径。
            3. 程序会自己打开 https://www.tianyancha.com/relation 网站，但需要用户自行登录，预设时间为30秒。
            4. 程序会自动读取用户选择的 Excel 文件中“Sheet1”中名为“公司列表”的列，对所有对象进行遍历，查询其两两之间的关系。
            5. 程序会下载公司的关系图，并将其整理放入一份名为“公司关联关系图”的 word 文档中，其位置为程序所处的文件夹。
            6. 如果您想了解更多信息，下载最新版本程序，或是获取程序源码，请访问程序的网站："""

    top = tk.Toplevel(root)
    top.title("使用说明")
    tk.Label(top, text=instructions, justify=tk.LEFT, wraplength=400).pack(padx=10, pady=10)

    link_label = tk.Label(top, text="https://github.com/worldtree43/Tianyancha_guanxi", 
                         fg="blue", cursor="hand2", justify=tk.LEFT)
    link_label.pack()
    link_label.bind("<Button-1>", lambda event: webbrowser.open("https://github.com/worldtree43/Tianyancha_guanxi"))

def show_disclaimer():
    disclaimer = """免责声明：
    本程序仅供学习参考使用，任何人或组织不得将本仓库的内容用于非法用途或侵犯他人合法权益。
    本程序所涉及的爬虫技术仅用于学习和研究，不得用于对其他平台进行大规模爬虫或其他非法行为。
    对于因使用本程序内容而引起的任何法律责任，本程序不承担任何责任。
    使用本程序即表示您同意本免责声明的所有条款和条件。"""

    top = tk.Toplevel(root)
    top.title("免责声明")
    tk.Label(top, text=disclaimer, justify=tk.LEFT).pack(padx=10, pady=10)

def get_companies(file_path, sheet_name, column_name):
    try:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        return df[column_name].tolist()
    except Exception as e:
        print(f"读取Excel文件出错: {str(e)}")
        return []

def wait_mask_disappear(driver, timeout=10):
    """等待遮罩层消失"""
    try:
        WebDriverWait(driver, timeout).until_not(
            EC.presence_of_element_located((By.CLASS_NAME, "_703b485d"))
        )
    except Exception:
        pass  # 没有遮罩层直接跳过

def select_company_suggestion(driver, company_name):
    """只在当前可见的ant-popover里找下拉建议，优先点击em内容等于公司名的，否则兜底点第一个"""
    try:
        # 等待当前可见的ant-popover出现
        WebDriverWait(driver, 10).until(
            EC.visibility_of_element_located(
                (By.XPATH, "//div[contains(@class, 'ant-popover') and not(contains(@class, 'ant-popover-hidden'))]")
            )
        )
        popovers = driver.find_elements(By.XPATH, "//div[contains(@class, 'ant-popover') and not(contains(@class, 'ant-popover-hidden'))]")
        found = False
        for popover in popovers:
            options = popover.find_elements(By.CLASS_NAME, "_703b485d")
            for option in options:
                try:
                    em_elem = option.find_element(By.TAG_NAME, "em")
                    em_text = em_elem.text.strip()
                    if em_text == company_name:
                        driver.execute_script("arguments[0].click();", option)
                        found = True
                        break
                except Exception:
                    continue
            if found:
                break
        # 没找到完全匹配，兜底点第一个
        if not found and popovers:
            options = popovers[0].find_elements(By.CLASS_NAME, "_703b485d")
            if options:
                driver.execute_script("arguments[0].click();", options[0])
        time.sleep(0.5)
    except Exception as e:
        print("未找到下拉建议，直接跳过:", e)

def tianyancha_relation_screenshot(companies, download_folder):
    if not companies:
        print("公司列表为空，请检查Excel文件")
        return

    chrome_options = Options()
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=1920x1080")

    driver = webdriver.Chrome(options=chrome_options)
    driver.maximize_window()

    try:
        url = 'https://www.tianyancha.com/relation'
        driver.get(url)
        input("请手动登录并进入关系图页面，页面加载完全后按回车继续...")

        wait = WebDriverWait(driver, 15)

        for i in range(len(companies)):
            for j in range(i+1, len(companies)):
                try:
                    # 定位所有输入框
                    inputs = wait.until(
                        EC.presence_of_all_elements_located(
                            (By.XPATH, '//input[@placeholder="请输入企业、人员名称"]')
                        )
                    )
                    if len(inputs) < 2:
                        raise Exception("未找到两个输入框，请检查页面结构！")

                    # 主体1
                    inputs[0].click()
                    inputs[0].clear()
                    time.sleep(0.3)
                    inputs[0].send_keys(companies[i])
                    time.sleep(0.5)
                    select_company_suggestion(driver, companies[i])

                    # 主体2
                    inputs[1].click()
                    inputs[1].clear()
                    time.sleep(0.3)
                    inputs[1].send_keys(companies[j])
                    time.sleep(0.5)
                    select_company_suggestion(driver, companies[j])

                    # 点击“开始分析”按钮
                    search_btn = wait.until(
                        EC.element_to_be_clickable(
                            (By.XPATH, '//button[.//span[text()="开始分析"]]')
                        )
                    )
                    wait_mask_disappear(driver)
                    driver.execute_script("arguments[0].click();", search_btn)

                    # 等待结果加载
                    time.sleep(random.uniform(7, 10))

                    # 判断是否有“暂无数据”或“无关联”提示
                    if len(driver.find_elements(By.XPATH, '//*[contains(text(),"暂无数据") or contains(text(),"无关联")]')) > 0:
                        print(f"{companies[i]} 和 {companies[j]} 无关联，跳过")
                        continue

                    # 点击“导出全部”按钮
                    export_btn = wait.until(
                        EC.element_to_be_clickable(
                            (By.XPATH, '//button[.//span[contains(text(),"导出全部")]]')
                        )
                    )
                    wait_mask_disappear(driver)
                    driver.execute_script("arguments[0].click();", export_btn)
                    time.sleep(1)  # 等待弹窗出现

                    # 新增：点击弹窗“确定”按钮
                    try:
                        confirm_btn = WebDriverWait(driver, 10).until(
                            EC.element_to_be_clickable(
                                (By.XPATH, '//button[.//span[text()="确定"]]')
                            )
                        )
                        driver.execute_script("arguments[0].click();", confirm_btn)
                        time.sleep(1)
                    except Exception as e:
                        print("未找到导出确认弹窗或确定按钮，跳过:", e)

                    time.sleep(random.uniform(2, 5))

                except Exception as e:
                    print(f"处理 {companies[i]} 和 {companies[j]} 时出错: {str(e)}")
                    traceback.print_exc()
                    driver.save_screenshot("error_page.png")
                    with open("error_page.html", "w", encoding="utf-8") as f:
                        f.write(driver.page_source)
                    continue

    except Exception as e:
        print(f"浏览器操作出错: {str(e)}")
    finally:
        driver.quit()

def create_word_document(images_info, downloads_folder):
    doc = Document()
    for image_name, text in images_info:
        try:
            image_path = os.path.join(downloads_folder, image_name)
            if os.path.exists(image_path):
                doc.add_paragraph().add_run().add_picture(image_path, width=Inches(6.0))
                doc.add_paragraph(text)
        except Exception as e:
            print(f"添加图片 {image_name} 到Word时出错: {str(e)}")
    
    doc.save("公司关联关系图.docx")
    print("Word文档已生成：公司关联关系图.docx")

def process_files():
    if not excel_file:
        print("请先选择Excel文件")
        return
    if not downloads_folder:
        print("请先选择下载文件夹")
        return

    sheet_name = 'Sheet1'
    company_column = '公司列表'
    
    try:
        companies = get_companies(excel_file, sheet_name, company_column)
        if not companies:
            print("未找到公司列表，请检查Excel文件格式")
            return

        tianyancha_relation_screenshot(companies, downloads_folder)

        images_info = []
        for i in range(len(companies)):
            for j in range(i + 1, len(companies)):
                file_name = f"查关系图谱-{companies[i]}&{companies[j]}-天眼查.png"
                text = f"{companies[i]} & {companies[j]} 关联关系\n"
                images_info.append((file_name, text))

        create_word_document(images_info, downloads_folder)
        print("处理完成！")
        
    except Exception as e:
        print(f"处理文件时出错: {str(e)}")

# 创建主界面
root = tk.Tk()
root.title("天眼查关联关系图片批量下载程序")
root.geometry("600x400")

# 创建框架
frame = tk.Frame(root, padx=20, pady=20)
frame.pack(expand=True, fill='both')

# 按钮框架
button_frame = tk.Frame(frame)
button_frame.pack(pady=10)

# 说明和免责声明按钮
instructions_btn = tk.Button(button_frame, text="使用说明", command=show_instructions)
instructions_btn.pack(side=tk.LEFT, padx=5)

disclaimer_btn = tk.Button(button_frame, text="免责声明", command=show_disclaimer)
disclaimer_btn.pack(side=tk.LEFT, padx=5)

# 文件选择部分
tk.Label(frame, text="1. 请选择包含公司列表的Excel文件：").pack(anchor='w', pady=(10, 0))
excel_btn = tk.Button(frame, text="选择Excel文件", command=choose_excel_file)
excel_btn.pack(pady=5)
excel_label = tk.Label(frame, text="未选择文件", wraplength=500, justify='left')
excel_label.pack(anchor='w')

# 文件夹选择部分
tk.Label(frame, text="\n2. 请选择下载文件夹：").pack(anchor='w', pady=(10, 0))
folder_btn = tk.Button(frame, text="选择下载文件夹", command=choose_downloads_folder)
folder_btn.pack(pady=5)
folder_label = tk.Label(frame, text="未选择文件夹", wraplength=500, justify='left')
folder_label.pack(anchor='w')

# 开始处理按钮
process_btn = tk.Button(frame, text="开始处理", command=process_files, bg="#4CAF50", fg="white", padx=20, pady=10)
process_btn.pack(pady=20)

# 状态标签
status_label = tk.Label(frame, text="", fg="green")
status_label.pack(pady=10)

# 启动主循环
root.mainloop()
